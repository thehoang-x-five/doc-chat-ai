"""
Document conversion task.
Converts documents to various output formats (TXT, MD, JSON, PDF, DOCX).
"""
import json
import tempfile
from pathlib import Path
from uuid import UUID
from datetime import timedelta
from celery import shared_task
from celery.utils.log import get_task_logger

logger = get_task_logger(__name__)


# Supported output formats
SUPPORTED_FORMATS = ["txt", "md", "json", "pdf", "docx", "html", "rtf"]


@shared_task(
    bind=True,
    name="app.queue.tasks.convert.process_convert",
    queue="convert",
    max_retries=3,
    default_retry_delay=60,
)
def process_convert(self, job_id: str, document_version_id: str, config: dict = None):
    """
    Convert a document to specified format.
    
    Args:
        job_id: Job UUID string
        document_version_id: DocumentVersion UUID string
        config: Configuration dict with:
            - output_format: Target format (txt, md, json, pdf, docx)
            - include_metadata: Include document metadata in output
    """
    from app.services.job_service import JobService
    from app.db.session import SessionLocal
    from app.db.models import DocumentVersion, Document, Chunk
    from app.storage.object_store import ObjectStore
    from sqlalchemy import select
    
    config = config or {}
    output_format = config.get("output_format", "txt").lower()
    include_metadata = config.get("include_metadata", True)
    
    if output_format not in SUPPORTED_FORMATS:
        raise ValueError(f"Unsupported format: {output_format}. Supported: {SUPPORTED_FORMATS}")
    
    logger.info(f"Starting CONVERT job {job_id} to {output_format}")
    
    try:
        with SessionLocal() as session:
            job_service = JobService(session)
            storage = ObjectStore()
            
            # Update job status to RUNNING
            job_service.update_status_sync(
                UUID(job_id),
                status="RUNNING",
                step="loading_document",
                progress=10
            )
            session.commit()
            
            # Load document version
            result = session.execute(
                select(DocumentVersion)
                .where(DocumentVersion.id == UUID(document_version_id))
            )
            doc_version = result.scalar_one_or_none()
            
            if not doc_version:
                raise ValueError(f"DocumentVersion {document_version_id} not found")
            
            # Load document
            result = session.execute(
                select(Document).where(Document.id == doc_version.document_id)
            )
            document = result.scalar_one_or_none()
            
            if not document:
                raise ValueError(f"Document not found")
            
            job_service.update_status_sync(
                UUID(job_id),
                step="loading_content",
                progress=20
            )
            session.commit()
            
            # Load chunks for this version
            result = session.execute(
                select(Chunk)
                .where(Chunk.document_version_id == doc_version.id)
                .order_by(Chunk.chunk_index)
            )
            chunks = list(result.scalars().all())
            
            # Get OCR output if available
            ocr_content = None
            if doc_version.ocr_output_key:
                try:
                    ocr_data = storage.download(doc_version.ocr_output_key)
                    ocr_content = json.loads(ocr_data.decode('utf-8'))
                except Exception as e:
                    logger.warning(f"Could not load OCR output: {e}")
            
            job_service.update_status_sync(
                UUID(job_id),
                step="converting",
                progress=40
            )
            session.commit()
            
            # Convert based on format
            if output_format == "txt":
                output_content, content_type = _convert_to_txt(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "md":
                output_content, content_type = _convert_to_markdown(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "json":
                output_content, content_type = _convert_to_json(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "pdf":
                output_content, content_type = _convert_to_pdf(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "docx":
                output_content, content_type = _convert_to_docx(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "html":
                output_content, content_type = _convert_to_html(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            elif output_format == "rtf":
                output_content, content_type = _convert_to_rtf(
                    document, doc_version, chunks, ocr_content, include_metadata
                )
            else:
                raise ValueError(f"Unsupported format: {output_format}")
            
            job_service.update_status_sync(
                UUID(job_id),
                step="storing",
                progress=70
            )
            session.commit()
            
            # Store converted file
            output_filename = f"{document.title}.{output_format}"
            output_key = storage.generate_key(
                str(document.workspace_id),
                str(document.id),
                output_filename,
                prefix="converted"
            )
            
            if isinstance(output_content, str):
                output_bytes = output_content.encode('utf-8')
            else:
                output_bytes = output_content
            
            storage.upload(output_key, output_bytes, content_type=content_type)
            
            job_service.update_status_sync(
                UUID(job_id),
                step="generating_url",
                progress=90
            )
            session.commit()
            
            # Generate presigned URL
            presigned_url = storage.get_presigned_url(
                output_key,
                expires=timedelta(hours=24)
            )
            
            # Mark as done with result
            job_service.update_status_sync(
                UUID(job_id),
                status="DONE",
                step="completed",
                progress=100,
                result_json={
                    "output_key": output_key,
                    "output_format": output_format,
                    "download_url": presigned_url,
                    "size_bytes": len(output_bytes),
                }
            )
            session.commit()
            
            logger.info(f"CONVERT job {job_id} completed: {output_key}")
            return {
                "status": "success",
                "job_id": job_id,
                "output_key": output_key,
                "download_url": presigned_url,
            }
            
    except Exception as e:
        logger.error(f"CONVERT job {job_id} failed: {e}")
        
        try:
            with SessionLocal() as session:
                job_service = JobService(session)
                job_service.update_status_sync(
                    UUID(job_id),
                    status="ERROR",
                    error_message=str(e)
                )
                session.commit()
        except Exception:
            pass
        
        raise self.retry(exc=e)


def _convert_to_txt(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to plain text format."""
    lines = []
    
    if include_metadata:
        lines.append(f"Title: {document.title}")
        lines.append(f"Type: {document.doc_type}")
        lines.append(f"Version: {doc_version.version}")
        lines.append(f"Created: {document.created_at}")
        lines.append("")
        lines.append("=" * 50)
        lines.append("")
    
    # Use OCR text if available
    if ocr_content and "text" in ocr_content:
        lines.append(ocr_content["text"])
    elif chunks:
        for chunk in chunks:
            lines.append(chunk.content)
            lines.append("")
    
    return "\n".join(lines), "text/plain"


def _convert_to_markdown(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to Markdown format."""
    lines = []
    
    lines.append(f"# {document.title}")
    lines.append("")
    
    if include_metadata:
        lines.append("## Metadata")
        lines.append("")
        lines.append(f"- **Type:** {document.doc_type}")
        lines.append(f"- **Version:** {doc_version.version}")
        lines.append(f"- **Created:** {document.created_at}")
        if document.tags:
            lines.append(f"- **Tags:** {', '.join(document.tags)}")
        lines.append("")
        lines.append("---")
        lines.append("")
    
    lines.append("## Content")
    lines.append("")
    
    # Use OCR markdown if available
    if ocr_content and "markdown" in ocr_content:
        lines.append(ocr_content["markdown"])
    elif ocr_content and "text" in ocr_content:
        lines.append(ocr_content["text"])
    elif chunks:
        current_page = None
        for chunk in chunks:
            if chunk.page_start and chunk.page_start != current_page:
                current_page = chunk.page_start
                lines.append(f"\n### Page {current_page}\n")
            lines.append(chunk.content)
            lines.append("")
    
    return "\n".join(lines), "text/markdown"


def _convert_to_json(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to JSON format."""
    output = {
        "document": {
            "id": str(document.id),
            "title": document.title,
            "type": document.doc_type,
            "source": document.source,
            "tags": document.tags or [],
            "status": document.status.value if hasattr(document.status, 'value') else str(document.status),
            "created_at": document.created_at.isoformat() if document.created_at else None,
        },
        "version": {
            "id": str(doc_version.id),
            "version": doc_version.version,
            "mime_type": doc_version.mime_type,
            "size_bytes": doc_version.size_bytes,
            "checksum": doc_version.checksum_sha256,
        },
        "content": {
            "chunks": [],
            "full_text": "",
        }
    }
    
    # Add OCR content if available
    if ocr_content:
        output["ocr"] = ocr_content
        if "text" in ocr_content:
            output["content"]["full_text"] = ocr_content["text"]
    
    # Add chunks
    for chunk in chunks:
        output["content"]["chunks"].append({
            "index": chunk.chunk_index,
            "content": chunk.content,
            "page_start": chunk.page_start,
            "page_end": chunk.page_end,
            "section": chunk.section_title,
            "token_count": chunk.token_count,
        })
    
    # Build full text from chunks if not from OCR
    if not output["content"]["full_text"] and chunks:
        output["content"]["full_text"] = "\n\n".join(c.content for c in chunks)
    
    return json.dumps(output, indent=2, ensure_ascii=False), "application/json"


def _convert_to_pdf(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to PDF format using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import io
    except ImportError:
        # Fallback: return text with PDF header
        logger.warning("reportlab not installed, returning text-based PDF")
        text_content, _ = _convert_to_txt(document, doc_version, chunks, ocr_content, include_metadata)
        return text_content.encode('utf-8'), "application/pdf"
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=0.5*inch, bottomMargin=0.5*inch)
    
    styles = getSampleStyleSheet()
    title_style = styles['Heading1']
    normal_style = styles['Normal']
    
    story = []
    
    # Title
    story.append(Paragraph(document.title, title_style))
    story.append(Spacer(1, 12))
    
    if include_metadata:
        meta_style = ParagraphStyle('meta', parent=normal_style, fontSize=10, textColor='gray')
        story.append(Paragraph(f"Type: {document.doc_type} | Version: {doc_version.version}", meta_style))
        story.append(Spacer(1, 24))
    
    # Content
    content_text = ""
    if ocr_content and "text" in ocr_content:
        content_text = ocr_content["text"]
    elif chunks:
        content_text = "\n\n".join(c.content for c in chunks)
    
    # Split into paragraphs
    for para in content_text.split("\n\n"):
        if para.strip():
            # Escape special characters for reportlab
            safe_para = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            try:
                story.append(Paragraph(safe_para, normal_style))
                story.append(Spacer(1, 6))
            except Exception:
                # If paragraph fails, add as plain text
                story.append(Paragraph(safe_para[:500] + "...", normal_style))
    
    doc.build(story)
    buffer.seek(0)
    
    return buffer.read(), "application/pdf"


def _convert_to_docx(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to DOCX format using python-docx."""
    try:
        from docx import Document as DocxDocument
        from docx.shared import Inches, Pt
        import io
    except ImportError:
        # Fallback: return text
        logger.warning("python-docx not installed, returning text")
        text_content, _ = _convert_to_txt(document, doc_version, chunks, ocr_content, include_metadata)
        return text_content.encode('utf-8'), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    docx = DocxDocument()
    
    # Title
    docx.add_heading(document.title, 0)
    
    if include_metadata:
        meta_para = docx.add_paragraph()
        meta_para.add_run(f"Type: {document.doc_type} | Version: {doc_version.version}").italic = True
        docx.add_paragraph()
    
    # Content
    if ocr_content and "text" in ocr_content:
        content_text = ocr_content["text"]
    elif chunks:
        content_text = "\n\n".join(c.content for c in chunks)
    else:
        content_text = ""
    
    # Add content paragraphs
    for para in content_text.split("\n\n"):
        if para.strip():
            docx.add_paragraph(para)
    
    # Save to buffer
    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    
    return buffer.read(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"



def _convert_to_html(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to HTML format."""
    html_parts = []
    
    # HTML header
    html_parts.append("<!DOCTYPE html>")
    html_parts.append("<html lang='vi'>")
    html_parts.append("<head>")
    html_parts.append(f"  <meta charset='UTF-8'>")
    html_parts.append(f"  <meta name='viewport' content='width=device-width, initial-scale=1.0'>")
    html_parts.append(f"  <title>{_escape_html(document.title)}</title>")
    html_parts.append("  <style>")
    html_parts.append("    body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; line-height: 1.6; }")
    html_parts.append("    h1 { color: #333; border-bottom: 2px solid #007bff; padding-bottom: 10px; }")
    html_parts.append("    .metadata { background: #f8f9fa; padding: 15px; border-radius: 8px; margin-bottom: 20px; }")
    html_parts.append("    .metadata dt { font-weight: bold; color: #666; }")
    html_parts.append("    .metadata dd { margin-left: 0; margin-bottom: 8px; }")
    html_parts.append("    .content { text-align: justify; }")
    html_parts.append("    .page-break { border-top: 1px dashed #ccc; margin: 20px 0; padding-top: 10px; color: #999; font-size: 0.9em; }")
    html_parts.append("  </style>")
    html_parts.append("</head>")
    html_parts.append("<body>")
    
    # Title
    html_parts.append(f"  <h1>{_escape_html(document.title)}</h1>")
    
    # Metadata
    if include_metadata:
        html_parts.append("  <div class='metadata'>")
        html_parts.append("    <dl>")
        html_parts.append(f"      <dt>Type</dt><dd>{_escape_html(document.doc_type)}</dd>")
        html_parts.append(f"      <dt>Version</dt><dd>{doc_version.version}</dd>")
        html_parts.append(f"      <dt>Created</dt><dd>{document.created_at}</dd>")
        if document.tags:
            html_parts.append(f"      <dt>Tags</dt><dd>{_escape_html(', '.join(document.tags))}</dd>")
        html_parts.append("    </dl>")
        html_parts.append("  </div>")
    
    # Content
    html_parts.append("  <div class='content'>")
    
    if ocr_content and "text" in ocr_content:
        content_text = ocr_content["text"]
        for para in content_text.split("\n\n"):
            if para.strip():
                html_parts.append(f"    <p>{_escape_html(para)}</p>")
    elif chunks:
        current_page = None
        for chunk in chunks:
            if chunk.page_start and chunk.page_start != current_page:
                current_page = chunk.page_start
                html_parts.append(f"    <div class='page-break'>Page {current_page}</div>")
            html_parts.append(f"    <p>{_escape_html(chunk.content)}</p>")
    
    html_parts.append("  </div>")
    html_parts.append("</body>")
    html_parts.append("</html>")
    
    return "\n".join(html_parts), "text/html"


def _convert_to_rtf(document, doc_version, chunks, ocr_content, include_metadata) -> tuple:
    """Convert to RTF (Rich Text Format)."""
    rtf_parts = []
    
    # RTF header
    rtf_parts.append(r"{\rtf1\ansi\deff0")
    rtf_parts.append(r"{\fonttbl{\f0 Arial;}}")
    rtf_parts.append(r"{\colortbl;\red0\green0\blue0;\red128\green128\blue128;}")
    
    # Title
    rtf_parts.append(r"\pard\qc\b\fs36 " + _escape_rtf(document.title) + r"\b0\par")
    rtf_parts.append(r"\par")
    
    # Metadata
    if include_metadata:
        rtf_parts.append(r"\pard\cf2\fs20")
        rtf_parts.append(r"Type: " + _escape_rtf(document.doc_type) + r" | ")
        rtf_parts.append(r"Version: " + str(doc_version.version) + r" | ")
        rtf_parts.append(r"Created: " + str(document.created_at))
        rtf_parts.append(r"\cf1\par\par")
    
    # Content
    rtf_parts.append(r"\pard\fs24")
    
    if ocr_content and "text" in ocr_content:
        content_text = ocr_content["text"]
        for para in content_text.split("\n\n"):
            if para.strip():
                rtf_parts.append(_escape_rtf(para) + r"\par\par")
    elif chunks:
        current_page = None
        for chunk in chunks:
            if chunk.page_start and chunk.page_start != current_page:
                current_page = chunk.page_start
                rtf_parts.append(r"\par\cf2\i Page " + str(current_page) + r"\i0\cf1\par\par")
            rtf_parts.append(_escape_rtf(chunk.content) + r"\par\par")
    
    # RTF footer
    rtf_parts.append("}")
    
    return "".join(rtf_parts), "application/rtf"


def _escape_html(text: str) -> str:
    """Escape HTML special characters."""
    if not text:
        return ""
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
        .replace("'", "&#39;")
    )


def _escape_rtf(text: str) -> str:
    """Escape RTF special characters."""
    if not text:
        return ""
    result = str(text)
    # Escape backslash first
    result = result.replace("\\", "\\\\")
    # Escape braces
    result = result.replace("{", "\\{")
    result = result.replace("}", "\\}")
    # Handle newlines
    result = result.replace("\n", "\\par ")
    return result
